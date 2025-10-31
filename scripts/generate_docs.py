#!/usr/bin/env python3
"""
Generate project documentation in Word format
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def add_heading(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return para

def add_bullet(doc, text):
    """Add a bullet point"""
    return doc.add_paragraph(text, style='List Bullet')

def add_code_block(doc, code):
    """Add a code block with monospace font"""
    para = doc.add_paragraph(code)
    para.style = 'Intense Quote'
    run = para.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return para

def create_infrastructure_doc():
    """Create the Cloud Infrastructure Report"""
    doc = Document()

    # Title
    title = doc.add_heading('MedicalDocAI - Cloud Infrastructure Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph('A Beginner-Friendly Guide to AWS Services and Deployment')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph('This document explains all the cloud services used in the MedicalDocAI project and how they were set up. Even if you\'re new to cloud computing, you\'ll understand what each service does and why we use it.')

    # Table of Contents
    add_heading(doc, 'Table of Contents', 1)
    add_bullet(doc, 'Overview')
    add_bullet(doc, 'AWS Services Used')
    add_bullet(doc, 'Infrastructure Setup (eksctl)')
    add_bullet(doc, 'GitHub Actions CI/CD')
    add_bullet(doc, 'How Everything Works Together')
    add_bullet(doc, 'Cost Optimization')

    doc.add_page_break()

    # Overview Section
    add_heading(doc, '1. Overview', 1)
    add_paragraph(doc, 'MedicalDocAI is deployed on Amazon Web Services (AWS), which is a cloud platform that provides servers, storage, and other computing resources over the internet. Instead of buying and maintaining physical servers, we rent cloud resources that automatically scale based on our needs.')

    doc.add_paragraph()
    add_paragraph(doc, 'Key Benefits of Cloud Deployment:', bold=True)
    add_bullet(doc, 'Scalability: Automatically handle more users without manual intervention')
    add_bullet(doc, 'Reliability: AWS provides 99.99% uptime with built-in redundancy')
    add_bullet(doc, 'Cost-effective: Only pay for what you use')
    add_bullet(doc, 'Security: Enterprise-grade security with IAM, encryption, and monitoring')

    doc.add_page_break()

    # AWS Services Section
    add_heading(doc, '2. AWS Services Used', 1)

    # EKS
    add_heading(doc, '2.1 Amazon EKS (Elastic Kubernetes Service)', 2)
    add_paragraph(doc, 'What it is: A managed Kubernetes service that runs containerized applications.')
    doc.add_paragraph()
    add_paragraph(doc, 'Think of it as:', italic=True)
    doc.add_paragraph('Imagine you have multiple apps (frontend, backend, database) that need to work together. Kubernetes is like a smart manager that:\n• Decides which server (called a "node") should run each app\n• Automatically restarts apps if they crash\n• Distributes traffic across multiple copies of your app for better performance\n• Scales up or down based on demand')

    doc.add_paragraph()
    add_paragraph(doc, 'Our Configuration:', bold=True)
    add_bullet(doc, 'Cluster Name: med-doc-dev')
    add_bullet(doc, 'Region: ap-south-1 (Mumbai, India)')
    add_bullet(doc, 'Kubernetes Version: 1.30')
    add_bullet(doc, 'Node Type: t3.small spot instances (cost-optimized)')
    add_bullet(doc, 'Node Count: 2 nodes (can scale 1-3)')

    doc.add_paragraph()
    add_paragraph(doc, 'Why EKS?', bold=True)
    add_bullet(doc, 'Automatic updates and security patches')
    add_bullet(doc, 'Built-in monitoring with CloudWatch')
    add_bullet(doc, 'Easy integration with other AWS services')
    add_bullet(doc, 'High availability across multiple data centers')

    doc.add_page_break()

    # ECR
    add_heading(doc, '2.2 Amazon ECR (Elastic Container Registry)', 2)
    add_paragraph(doc, 'What it is: A secure place to store Docker container images.')
    doc.add_paragraph()
    add_paragraph(doc, 'Think of it as:', italic=True)
    doc.add_paragraph('Like a photo album, but instead of photos, it stores "snapshots" of your application. Each snapshot (called a container image) contains:\n• Your application code\n• All dependencies and libraries\n• The operating system environment\n• Configuration files')

    doc.add_paragraph()
    add_paragraph(doc, 'Our Repositories:', bold=True)
    add_bullet(doc, 'Backend Repository: med-doc-backend (FastAPI application)')
    add_bullet(doc, 'Frontend Repository: med-doc-frontend (React application)')

    doc.add_paragraph()
    add_paragraph(doc, 'How it works:', bold=True)
    doc.add_paragraph('1. Developer pushes code to GitHub\n2. GitHub Actions builds a Docker image (like creating a snapshot)\n3. The image is pushed to ECR (uploaded to the album)\n4. EKS pulls the image from ECR and runs it (displays the photo)')

    doc.add_page_break()

    # S3
    add_heading(doc, '2.3 Amazon S3 (Simple Storage Service)', 2)
    add_paragraph(doc, 'What it is: Cloud storage for files and data.')
    doc.add_paragraph()
    add_paragraph(doc, 'Think of it as:', italic=True)
    doc.add_paragraph('Like Google Drive or Dropbox, but designed for applications. It stores files in "buckets" (folders) that can be accessed from anywhere.')

    doc.add_paragraph()
    add_paragraph(doc, 'Our Usage:', bold=True)
    add_bullet(doc, 'Bucket Name: med-docs-dev')
    add_bullet(doc, 'Purpose: Store evaluation results')
    add_bullet(doc, 'Location: eval-results/latest.json (most recent eval)')
    add_bullet(doc, 'Location: eval-results/<commit-sha>.json (historical results)')

    doc.add_paragraph()
    add_paragraph(doc, 'Why S3 for eval results?', bold=True)
    add_bullet(doc, 'Running evaluations is expensive (uses Claude API)')
    add_bullet(doc, 'We pre-compute results during deployment and cache them')
    add_bullet(doc, 'Users get instant results without re-running tests')
    add_bullet(doc, 'We maintain a history of all past evaluations')

    doc.add_page_break()

    # IAM
    add_heading(doc, '2.4 AWS IAM (Identity and Access Management)', 2)
    add_paragraph(doc, 'What it is: Security system that controls who can access what.')
    doc.add_paragraph()
    add_paragraph(doc, 'Think of it as:', italic=True)
    doc.add_paragraph('Like a security guard with a list of permissions. It ensures:\n• GitHub Actions can deploy to EKS\n• Backend pods can read/write to S3\n• No unauthorized access to sensitive resources')

    doc.add_paragraph()
    add_paragraph(doc, 'Our IAM Setup:', bold=True)

    doc.add_paragraph()
    add_paragraph(doc, 'a) GitHub Actions Role:', bold=True)
    add_bullet(doc, 'Allows GitHub to authenticate with AWS')
    add_bullet(doc, 'Permissions: Push images to ECR, deploy to EKS')
    add_bullet(doc, 'Uses OpenID Connect (OIDC) for secure authentication')

    doc.add_paragraph()
    add_paragraph(doc, 'b) Backend Service Account (api-sa):', bold=True)
    add_bullet(doc, 'Allows backend pods to access AWS services')
    add_bullet(doc, 'Permissions: Read/write S3 bucket for eval results')
    add_bullet(doc, 'Role ARN: arn:aws:iam::503837496832:role/med-doc-api-role')

    doc.add_page_break()

    # Load Balancer
    add_heading(doc, '2.5 AWS Load Balancer', 2)
    add_paragraph(doc, 'What it is: Distributes incoming traffic across multiple servers.')
    doc.add_paragraph()
    add_paragraph(doc, 'Think of it as:', italic=True)
    doc.add_paragraph('Like a receptionist directing visitors to different offices. When users visit your website:\n• The load balancer receives the request\n• It picks a healthy backend server\n• It forwards the request to that server\n• If a server crashes, it automatically routes traffic elsewhere')

    doc.add_paragraph()
    add_paragraph(doc, 'Our Configuration:', bold=True)
    add_bullet(doc, 'Type: Classic Load Balancer')
    add_bullet(doc, 'Idle Timeout: 5 minutes (for long-running eval requests)')
    add_bullet(doc, 'Health Checks: Pings /health endpoint every 10 seconds')
    add_bullet(doc, 'Public URL: http://ad28ddbf3abfb43588a27433f489cfb0-826398851.ap-south-1.elb.amazonaws.com')

    doc.add_page_break()

    # CloudWatch
    add_heading(doc, '2.6 AWS CloudWatch', 2)
    add_paragraph(doc, 'What it is: Monitoring and logging service.')
    doc.add_paragraph()
    add_paragraph(doc, 'Think of it as:', italic=True)
    doc.add_paragraph('Like a security camera system that records everything:\n• Application logs (errors, info messages)\n• Performance metrics (CPU, memory usage)\n• API activity (who accessed what and when)')

    doc.add_paragraph()
    add_paragraph(doc, 'What we monitor:', bold=True)
    add_bullet(doc, 'EKS control plane logs (API server, scheduler, controller)')
    add_bullet(doc, 'Pod logs (backend and frontend application logs)')
    add_bullet(doc, 'Deployment events (successful/failed deployments)')

    doc.add_page_break()

    # Infrastructure Setup
    add_heading(doc, '3. Infrastructure Setup with eksctl', 1)

    add_paragraph(doc, 'We use eksctl (an official AWS tool) to create and manage our Kubernetes cluster. eksctl automatically generates CloudFormation templates behind the scenes.')

    doc.add_paragraph()
    add_heading(doc, 'What is CloudFormation?', 2)
    add_paragraph(doc, 'CloudFormation is AWS\'s infrastructure-as-code service. Instead of manually clicking buttons in the AWS console, you write a YAML file describing what you want, and AWS creates it automatically.')

    doc.add_paragraph()
    add_paragraph(doc, 'Think of it as:', italic=True)
    doc.add_paragraph('Like IKEA furniture instructions. You describe what you want (a chair), and the instructions (CloudFormation) tell AWS exactly how to build it step-by-step.')

    doc.add_paragraph()
    add_heading(doc, 'Our Setup Process', 2)

    add_paragraph(doc, 'Step 1: Create cluster configuration (cloud/eksctl/cluster.config.yaml)')
    add_code_block(doc, '''apiVersion: eksctl.io/v1alpha5
kind: ClusterConfig

metadata:
  name: med-doc-dev
  region: ap-south-1
  version: "1.30"

managedNodeGroups:
  - name: ng-spot
    instanceTypes: ["t3.small"]
    desiredCapacity: 2
    minSize: 1
    maxSize: 3
    spot: true

iam:
  withOIDC: true  # Enables service accounts
''')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 2: Create the cluster')
    add_code_block(doc, 'eksctl create cluster -f cloud/eksctl/cluster.config.yaml')

    doc.add_paragraph()
    add_paragraph(doc, 'What happens behind the scenes:', bold=True)
    add_bullet(doc, 'eksctl generates CloudFormation templates')
    add_bullet(doc, 'CloudFormation creates: VPC, subnets, security groups, EKS cluster, node groups')
    add_bullet(doc, 'Takes about 15-20 minutes to complete')
    add_bullet(doc, 'You can view progress in AWS Console > CloudFormation')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 3: Set up IAM roles')
    doc.add_paragraph('We manually created IAM roles using JSON policies in cloud/iam/ folder:')
    add_bullet(doc, 'github-actions-policy.json - Allows GitHub to deploy')
    add_bullet(doc, 's3-access-policy.json - Allows backend to access S3')
    add_bullet(doc, 'pod-trust-policy.json - Allows pods to assume IAM roles')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 4: Deploy applications')
    add_code_block(doc, '''# Apply Kubernetes manifests
kubectl apply -f k8s/base/iam-serviceaccount.yaml
kubectl apply -f k8s/base/api-deployment.yaml
kubectl apply -f k8s/base/web-deployment.yaml
''')

    doc.add_page_break()

    # CI/CD Section
    add_heading(doc, '4. GitHub Actions CI/CD Pipeline', 1)

    add_paragraph(doc, 'CI/CD stands for Continuous Integration / Continuous Deployment. It automates the process of building, testing, and deploying code.')

    doc.add_paragraph()
    add_paragraph(doc, 'Think of it as:', italic=True)
    doc.add_paragraph('Like a robot assistant that:\n• Watches your GitHub repository\n• When you push code, it automatically builds and tests it\n• If tests pass, it deploys to production\n• No manual steps required!')

    doc.add_paragraph()
    add_heading(doc, 'Our Pipeline (.github/workflows/ci-cd.yml)', 2)

    add_paragraph(doc, 'Step 1: Build Images', bold=True)
    add_bullet(doc, 'Triggered when code is pushed to main branch')
    add_bullet(doc, 'Builds Docker images for backend and frontend')
    add_bullet(doc, 'Pushes images to ECR with commit SHA tag')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 2: Deploy to Kubernetes', bold=True)
    add_bullet(doc, 'Updates EKS deployments with new image tags')
    add_bullet(doc, 'Waits for pods to become ready (health checks)')
    add_bullet(doc, 'Performs rolling update (zero downtime)')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 3: Run Evaluation Gate', bold=True)
    add_bullet(doc, 'Tests critical API endpoints (classify, extract-codes, summarize)')
    add_bullet(doc, 'If tests fail, deployment is marked as failed')
    add_bullet(doc, 'Ensures production is always working')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 4: Generate and Store Eval Results', bold=True)
    add_bullet(doc, 'Runs comprehensive evaluation with 50+ test cases')
    add_bullet(doc, 'Uploads results to S3 (eval-results/latest.json)')
    add_bullet(doc, 'Also saves with commit SHA for history')
    add_bullet(doc, 'Users fetch pre-computed results instead of re-running tests')

    doc.add_paragraph()
    add_heading(doc, 'Pipeline Duration', 2)
    add_bullet(doc, 'Build Phase: ~3-5 minutes')
    add_bullet(doc, 'Deploy Phase: ~2-3 minutes')
    add_bullet(doc, 'Evaluation Gate: ~1 minute')
    add_bullet(doc, 'Eval Generation: ~2-3 minutes')
    add_bullet(doc, 'Total: ~8-12 minutes per deployment')

    doc.add_page_break()

    # How Everything Works Together
    add_heading(doc, '5. How Everything Works Together', 1)

    add_paragraph(doc, 'Here\'s the complete flow from code to production:', bold=True)

    doc.add_paragraph()
    add_paragraph(doc, '1. Developer Pushes Code')
    doc.add_paragraph('   • Developer commits changes to hotfix/feature branch\n   • Code is reviewed and merged to main branch')

    doc.add_paragraph()
    add_paragraph(doc, '2. GitHub Actions Triggered')
    doc.add_paragraph('   • Workflow starts automatically\n   • Authenticates with AWS using OIDC (no stored credentials!)')

    doc.add_paragraph()
    add_paragraph(doc, '3. Docker Images Built')
    doc.add_paragraph('   • Backend image built from backend-fastapi/Dockerfile\n   • Frontend image built from Dockerfile.frontend\n   • Tagged with commit SHA (e.g., 5177d05c)')

    doc.add_paragraph()
    add_paragraph(doc, '4. Images Pushed to ECR')
    doc.add_paragraph('   • Images uploaded to ECR repositories\n   • Available for EKS to pull')

    doc.add_paragraph()
    add_paragraph(doc, '5. EKS Deployment Updated')
    doc.add_paragraph('   • kubectl updates deployment with new image tag\n   • Kubernetes performs rolling update:\n     - Starts new pods with new code\n     - Waits for health checks to pass\n     - Terminates old pods\n   • Zero downtime for users!')

    doc.add_paragraph()
    add_paragraph(doc, '6. Evaluation Tests Run')
    doc.add_paragraph('   • Smoke tests verify core functionality\n   • Full evaluation runs 50+ test cases\n   • Results uploaded to S3')

    doc.add_paragraph()
    add_paragraph(doc, '7. Production Live')
    doc.add_paragraph('   • New version accessible via load balancer\n   • Users see updated features\n   • Eval reports fetch from S3 cache')

    doc.add_page_break()

    # Cost Optimization
    add_heading(doc, '6. Cost Optimization Strategies', 1)

    add_paragraph(doc, 'We use several techniques to minimize AWS costs:', bold=True)

    doc.add_paragraph()
    add_heading(doc, '6.1 Spot Instances (70% savings)', 2)
    add_bullet(doc, 'Uses spare AWS capacity at steep discounts')
    add_bullet(doc, 'Can be interrupted (AWS gives 2-minute warning)')
    add_bullet(doc, 'Kubernetes automatically reschedules workloads')
    add_bullet(doc, 'Perfect for development environments')

    doc.add_paragraph()
    add_heading(doc, '6.2 Right-Sized Instances', 2)
    add_bullet(doc, 't3.small nodes (2 vCPU, 2GB RAM)')
    add_bullet(doc, 'Sufficient for current workload')
    add_bullet(doc, 'Can upgrade to t3.medium if needed')

    doc.add_paragraph()
    add_heading(doc, '6.3 Auto-Scaling', 2)
    add_bullet(doc, 'Node count: 1-3 (scales based on load)')
    add_bullet(doc, 'Pod replicas: 2 (can increase if traffic spikes)')
    add_bullet(doc, 'Only pay for what you use')

    doc.add_paragraph()
    add_heading(doc, '6.4 S3 Lifecycle Policies', 2)
    add_bullet(doc, 'Latest eval result cached indefinitely')
    add_bullet(doc, 'Historical results can be archived after 90 days')
    add_bullet(doc, 'Reduces storage costs over time')

    doc.add_paragraph()
    add_heading(doc, '6.5 Ephemeral SQLite Storage', 2)
    add_bullet(doc, 'Uses local disk instead of RDS database')
    add_bullet(doc, 'Saves $50-100/month')
    add_bullet(doc, 'Acceptable for development/demo purposes')

    doc.add_paragraph()
    add_heading(doc, 'Estimated Monthly Cost', 2)
    add_bullet(doc, 'EKS Control Plane: $72/month (fixed)')
    add_bullet(doc, 't3.small spot nodes (2x): ~$7/month')
    add_bullet(doc, 'Load Balancer: ~$18/month')
    add_bullet(doc, 'ECR Storage: ~$1/month')
    add_bullet(doc, 'S3 Storage: <$1/month')
    add_bullet(doc, 'Data Transfer: ~$5/month')
    doc.add_paragraph()
    add_paragraph(doc, 'Total: ~$104/month', bold=True)

    doc.add_paragraph()
    doc.add_paragraph('(Note: Claude API costs are separate and depend on usage)')

    doc.add_page_break()

    # Summary
    add_heading(doc, '7. Summary', 1)

    add_paragraph(doc, 'MedicalDocAI uses a modern, scalable cloud architecture:', bold=True)

    doc.add_paragraph()
    add_bullet(doc, 'EKS manages containerized applications with high availability')
    add_bullet(doc, 'ECR stores Docker images for reliable deployments')
    add_bullet(doc, 'S3 caches expensive evaluation results')
    add_bullet(doc, 'IAM ensures secure access control')
    add_bullet(doc, 'Load Balancer distributes traffic and provides public access')
    add_bullet(doc, 'GitHub Actions automates entire deployment pipeline')
    add_bullet(doc, 'Cost-optimized with spot instances and right-sizing')

    doc.add_paragraph()
    doc.add_paragraph('This setup provides enterprise-grade reliability and security while keeping costs under $110/month.')

    doc.add_paragraph()
    add_paragraph(doc, 'Key Takeaway:', bold=True)
    doc.add_paragraph('Instead of managing physical servers, we use cloud services that automatically handle scaling, security, monitoring, and deployment. The entire infrastructure is defined as code (Infrastructure-as-Code) using eksctl and Kubernetes manifests, making it reproducible and maintainable.')

    # Save document
    doc_path = os.path.join(os.path.dirname(__file__), '..', 'docs', 'Cloud_Infrastructure_Report.docx')
    os.makedirs(os.path.dirname(doc_path), exist_ok=True)
    doc.save(doc_path)
    print(f"[OK] Created: {doc_path}")
    return doc_path


def create_features_roadmap_doc():
    """Create the Features and Roadmap document"""
    doc = Document()

    # Title
    title = doc.add_heading('MedicalDocAI - Features & Roadmap', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph('Current Capabilities and Future Improvements')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph('This document describes all current features of the MedicalDocAI application and outlines planned improvements for future releases.')

    # Table of Contents
    add_heading(doc, 'Table of Contents', 1)
    add_bullet(doc, 'Current Features')
    add_bullet(doc, 'Technical Architecture')
    add_bullet(doc, 'Future Improvements')
    add_bullet(doc, 'Security Enhancements')
    add_bullet(doc, 'Scalability Roadmap')

    doc.add_page_break()

    # Current Features
    add_heading(doc, '1. Current Features', 1)

    add_heading(doc, '1.1 Document Classification (5 Types)', 2)
    add_paragraph(doc, 'Automatically identifies the type of medical document using Claude AI.')

    doc.add_paragraph()
    add_paragraph(doc, 'Supported Document Types:', bold=True)
    add_bullet(doc, 'COMPLETE BLOOD COUNT (CBC)')
    add_bullet(doc, 'BASIC METABOLIC PANEL (BMP)')
    add_bullet(doc, 'X-RAY')
    add_bullet(doc, 'CT (Computed Tomography)')
    add_bullet(doc, 'CLINICAL NOTE')

    doc.add_paragraph()
    add_paragraph(doc, 'How it works:', bold=True)
    add_bullet(doc, 'Upload a text file (.txt or .pdf)')
    add_bullet(doc, 'AI analyzes document content')
    add_bullet(doc, 'Returns classification with confidence score (0-1)')
    add_bullet(doc, 'Provides rationale and evidence for the classification')

    doc.add_paragraph()
    add_paragraph(doc, 'Example Output:')
    add_code_block(doc, '''{
  "document_type": "COMPLETE BLOOD COUNT",
  "confidence": 0.95,
  "rationale": "Document contains CBC lab values including WBC, Hemoglobin, and Platelets",
  "evidence": ["WBC 13.2 x10^3/µL", "Hgb 14.1 g/dL", "Platelets 250"]
}''')

    doc.add_page_break()

    # ICD-10 Code Extraction
    add_heading(doc, '1.2 ICD-10 Code Extraction with Evidence', 2)
    add_paragraph(doc, 'Extracts medical billing codes (ICD-10) from document text with supporting evidence.')

    doc.add_paragraph()
    add_paragraph(doc, 'Features:', bold=True)
    add_bullet(doc, 'Identifies multiple ICD-10 codes per document')
    add_bullet(doc, 'Provides confidence score for each code')
    add_bullet(doc, 'Shows evidence from document supporting each code')
    add_bullet(doc, 'Includes human-readable descriptions')

    doc.add_paragraph()
    add_paragraph(doc, 'Example Output:')
    add_code_block(doc, '''{
  "codes": [
    {
      "code": "D72.829",
      "description": "Elevated white blood cell count, unspecified",
      "confidence": 0.92,
      "evidence": ["WBC 13.2 x10^3/µL (elevated)", "leukocytosis"]
    }
  ]
}''')

    doc.add_paragraph()
    add_paragraph(doc, 'Why this matters:', bold=True)
    doc.add_paragraph('ICD-10 codes are required for medical billing and insurance claims. Automatic extraction saves time and reduces coding errors.')

    doc.add_page_break()

    # Clinical Summary
    add_heading(doc, '1.3 Provider-Facing Clinical Summary', 2)
    add_paragraph(doc, 'Generates a concise summary for healthcare providers highlighting key findings.')

    doc.add_paragraph()
    add_paragraph(doc, 'Features:', bold=True)
    add_bullet(doc, 'Brief 1-3 sentence summary')
    add_bullet(doc, 'Focuses on clinically significant findings')
    add_bullet(doc, 'References extracted ICD-10 codes')
    add_bullet(doc, 'Confidence score for summary quality')

    doc.add_paragraph()
    add_paragraph(doc, 'Example Output:')
    add_code_block(doc, '''{
  "summary": "Leukocytosis present with elevated WBC count. Hemoglobin and platelets within normal limits.",
  "confidence": 0.88,
  "evidence": ["WBC elevated", "normal Hgb", "normal platelets"]
}''')

    doc.add_page_break()

    # Document Type Hint
    add_heading(doc, '1.4 Document Type Hint Feature (NEW)', 2)
    add_paragraph(doc, 'Users can manually select document type to skip classification step.')

    doc.add_paragraph()
    add_paragraph(doc, 'Benefits:', bold=True)
    add_bullet(doc, 'Faster processing (skips classification)')
    add_bullet(doc, 'Higher accuracy when user knows document type')
    add_bullet(doc, 'Reduced API costs (one less Claude call)')
    add_bullet(doc, 'Better user experience for bulk uploads')

    doc.add_paragraph()
    add_paragraph(doc, 'UI Design:', bold=True)
    doc.add_paragraph('Dropdown selector with options:\n• AUTO — let AI classify (default)\n• Categorized by type:\n  📋 Labs: CBC, BMP\n  🏥 Imaging: X-Ray, CT\n  📄 Clinical Docs: Clinical Note')

    doc.add_paragraph()
    add_paragraph(doc, 'Technical Implementation:', bold=True)
    doc.add_paragraph('When user selects a type:\n• Backend receives document_type_hint parameter\n• Classification step is skipped\n• Confidence set to 1.0\n• Rationale: "User selected document type in UI; classification skipped."')

    doc.add_page_break()

    # Full Pipeline
    add_heading(doc, '1.5 Full Pipeline Processing', 2)
    add_paragraph(doc, 'All three steps (classify, extract codes, summarize) run automatically on document upload.')

    doc.add_paragraph()
    add_paragraph(doc, 'Pipeline Stages:', bold=True)
    doc.add_paragraph('1. Text Extraction (if PDF)')
    doc.add_paragraph('2. Classification (or skip if hint provided)')
    doc.add_paragraph('3. ICD-10 Code Extraction')
    doc.add_paragraph('4. Clinical Summary Generation')
    doc.add_paragraph('5. Store results in database')

    doc.add_paragraph()
    add_paragraph(doc, 'Processing Time:', bold=True)
    add_bullet(doc, 'With classification: ~8-12 seconds')
    add_bullet(doc, 'With hint: ~5-8 seconds')
    add_bullet(doc, 'Depends on document length and Claude API latency')

    doc.add_page_break()

    # Evaluation System
    add_heading(doc, '1.6 Evaluation System with S3 Caching', 2)
    add_paragraph(doc, 'Comprehensive testing framework to measure AI accuracy.')

    doc.add_paragraph()
    add_paragraph(doc, 'Features:', bold=True)
    add_bullet(doc, '50+ test cases across all document types')
    add_bullet(doc, 'Metrics: Precision, Recall, F1 Score, Summary Coverage')
    add_bullet(doc, 'Pre-computed during deployment (cached in S3)')
    add_bullet(doc, 'Historical tracking with commit SHA')

    doc.add_paragraph()
    add_paragraph(doc, 'Current Test Dataset:', bold=True)
    add_bullet(doc, '10 Complete Blood Count tests')
    add_bullet(doc, '10 Basic Metabolic Panel tests')
    add_bullet(doc, '10 X-Ray tests')
    add_bullet(doc, '10 CT tests')
    add_bullet(doc, '10 Clinical Note tests')

    doc.add_paragraph()
    add_paragraph(doc, 'Latest Performance (as of last deployment):', bold=True)
    add_bullet(doc, 'Precision: 59%')
    add_bullet(doc, 'Recall: 84%')
    add_bullet(doc, 'F1 Score: 69%')
    add_bullet(doc, 'Summary Coverage: 61%')

    doc.add_paragraph()
    add_paragraph(doc, 'How S3 Caching Works:', bold=True)
    doc.add_paragraph('1. GitHub Actions deploys new code\n2. CI/CD runs /eval/quick endpoint\n3. Results uploaded to s3://med-docs-dev/eval-results/latest.json\n4. Users click "Get Eval Report" → fetches from S3\n5. No expensive re-computation needed!')

    doc.add_paragraph()
    add_paragraph(doc, 'Local Eval Testing:', bold=True)
    doc.add_paragraph('Developers can test eval changes locally:')
    add_code_block(doc, '''# Run eval locally only
python scripts/run_local_eval.py

# Run eval and upload to S3 (update production)
python scripts/run_local_eval.py --upload-s3
''')

    doc.add_page_break()

    # RESTful API
    add_heading(doc, '1.7 RESTful API with FastAPI', 2)
    add_paragraph(doc, 'Well-documented API endpoints for all operations.')

    doc.add_paragraph()
    add_paragraph(doc, 'Available Endpoints:', bold=True)

    add_paragraph(doc, 'POST /classify')
    doc.add_paragraph('   Classify document type\n   Input: {"document_text": "..."}')

    doc.add_paragraph()
    add_paragraph(doc, 'POST /extract-codes')
    doc.add_paragraph('   Extract ICD-10 codes\n   Input: {"document_text": "...", "document_type": "..."}')

    doc.add_paragraph()
    add_paragraph(doc, 'POST /summarize')
    doc.add_paragraph('   Generate clinical summary\n   Input: {"document_text": "...", "document_type": "...", "codes": [...]}')

    doc.add_paragraph()
    add_paragraph(doc, 'POST /documents')
    doc.add_paragraph('   Upload and process document (full pipeline)\n   Input: FormData with file and optional document_type_hint')

    doc.add_paragraph()
    add_paragraph(doc, 'GET /documents/{id}')
    doc.add_paragraph('   Retrieve processed document results')

    doc.add_paragraph()
    add_paragraph(doc, 'GET /eval/quick')
    doc.add_paragraph('   Run comprehensive evaluation (50+ tests)')

    doc.add_paragraph()
    add_paragraph(doc, 'GET /eval/latest')
    doc.add_paragraph('   Fetch cached evaluation results from S3')

    doc.add_paragraph()
    add_paragraph(doc, 'GET /health')
    doc.add_paragraph('   Health check endpoint for monitoring')

    doc.add_paragraph()
    add_paragraph(doc, 'Interactive Documentation:', bold=True)
    doc.add_paragraph('FastAPI provides automatic Swagger UI:\n• Local: http://127.0.0.1:8000/docs\n• Production: <load-balancer-url>/docs')

    doc.add_page_break()

    # React Frontend
    add_heading(doc, '1.8 Modern React Frontend', 2)
    add_paragraph(doc, 'Intuitive web interface for document upload and results visualization.')

    doc.add_paragraph()
    add_paragraph(doc, 'Key Features:', bold=True)
    add_bullet(doc, 'Drag-and-drop file upload')
    add_bullet(doc, 'Document type selector with categories')
    add_bullet(doc, 'Real-time processing status')
    add_bullet(doc, 'Tabbed interface for results:')
    doc.add_paragraph('  • Classification Panel: Shows document type and confidence')
    doc.add_paragraph('  • ICD-10 Codes Panel: Interactive code cards with evidence')
    doc.add_paragraph('  • Summary Panel: Clinical summary with key findings')
    add_bullet(doc, 'Evaluation Report page with charts and metrics')
    add_bullet(doc, 'Responsive design (works on mobile/tablet)')

    doc.add_paragraph()
    add_paragraph(doc, 'Technology Stack:', bold=True)
    add_bullet(doc, 'React 18 with TypeScript')
    add_bullet(doc, 'Vite for fast development')
    add_bullet(doc, 'TanStack Query for data fetching')
    add_bullet(doc, 'Tailwind CSS for styling')
    add_bullet(doc, 'Lucide React for icons')

    doc.add_page_break()

    # Architecture
    add_heading(doc, '2. Technical Architecture', 1)

    add_heading(doc, '2.1 Backend Stack', 2)
    add_bullet(doc, 'FastAPI (Python 3.12) - Modern async web framework')
    add_bullet(doc, 'SQLAlchemy - ORM for database operations')
    add_bullet(doc, 'SQLite - Lightweight database (ephemeral in production)')
    add_bullet(doc, 'Anthropic SDK - Claude AI integration')
    add_bullet(doc, 'Boto3 - AWS SDK for S3 operations')
    add_bullet(doc, 'PyPDF2 - PDF text extraction')
    add_bullet(doc, 'Pydantic - Data validation and settings management')

    doc.add_paragraph()
    add_heading(doc, '2.2 Frontend Stack', 2)
    add_bullet(doc, 'React 18 - UI component library')
    add_bullet(doc, 'TypeScript - Type-safe JavaScript')
    add_bullet(doc, 'Vite - Build tool and dev server')
    add_bullet(doc, 'TanStack Query - Data fetching and caching')
    add_bullet(doc, 'Wouter - Lightweight routing')
    add_bullet(doc, 'Zod - Runtime type validation')
    add_bullet(doc, 'Tailwind CSS - Utility-first CSS framework')

    doc.add_paragraph()
    add_heading(doc, '2.3 Infrastructure', 2)
    add_bullet(doc, 'AWS EKS - Kubernetes cluster management')
    add_bullet(doc, 'AWS ECR - Docker image registry')
    add_bullet(doc, 'AWS S3 - Evaluation results storage')
    add_bullet(doc, 'AWS IAM - Access control and security')
    add_bullet(doc, 'AWS Load Balancer - Traffic distribution')
    add_bullet(doc, 'GitHub Actions - CI/CD pipeline')
    add_bullet(doc, 'Docker - Containerization')
    add_bullet(doc, 'Kubernetes - Container orchestration')

    doc.add_paragraph()
    add_heading(doc, '2.4 AI Integration', 2)
    add_bullet(doc, 'Claude Sonnet 4.5 - Primary AI model')
    add_bullet(doc, 'Prompt caching - Reduces latency and costs')
    add_bullet(doc, 'Structured outputs - Ensures consistent JSON responses')
    add_bullet(doc, 'Mock mode - For testing without API calls')

    doc.add_page_break()

    # Future Improvements
    add_heading(doc, '3. Future Improvements', 1)

    add_heading(doc, '3.1 Security Enhancements', 2)

    add_paragraph(doc, '3.1.1 Store Claude API Key in AWS Secrets Manager', bold=True)
    add_paragraph(doc, 'Current State: API key stored as Kubernetes secret')
    add_paragraph(doc, 'Improvement: Use AWS Secrets Manager')

    doc.add_paragraph()
    add_paragraph(doc, 'Benefits:', bold=True)
    add_bullet(doc, 'Automatic key rotation')
    add_bullet(doc, 'Audit logging (who accessed key and when)')
    add_bullet(doc, 'Fine-grained access control')
    add_bullet(doc, 'Encrypted at rest and in transit')
    add_bullet(doc, 'Integration with CloudTrail for compliance')

    doc.add_paragraph()
    add_paragraph(doc, 'Implementation Steps:', bold=True)
    doc.add_paragraph('''1. Store key in Secrets Manager:
   aws secretsmanager create-secret \\
     --name med-doc/anthropic-api-key \\
     --secret-string "sk-ant-..."

2. Update IAM role to allow reading secret

3. Modify backend to fetch from Secrets Manager:
   import boto3
   secrets = boto3.client('secretsmanager')
   api_key = secrets.get_secret_value(
       SecretId='med-doc/anthropic-api-key'
   )['SecretString']

4. Remove Kubernetes secret''')

    doc.add_paragraph()
    add_paragraph(doc, 'Estimated effort: 2-3 hours', italic=True)

    doc.add_page_break()

    add_paragraph(doc, '3.1.2 Implement User Authentication', bold=True)
    add_paragraph(doc, 'Current State: No authentication (open to public)')
    add_paragraph(doc, 'Improvement: Add user authentication system')

    doc.add_paragraph()
    add_paragraph(doc, 'Options:', bold=True)
    add_bullet(doc, 'AWS Cognito - Managed authentication service')
    add_bullet(doc, 'OAuth 2.0 - Google/Microsoft login integration')
    add_bullet(doc, 'JWT tokens - Stateless authentication')

    doc.add_paragraph()
    add_paragraph(doc, 'Features to Add:', bold=True)
    add_bullet(doc, 'User registration and login')
    add_bullet(doc, 'Role-based access control (admin, doctor, viewer)')
    add_bullet(doc, 'Document access permissions')
    add_bullet(doc, 'Audit trail of user actions')

    doc.add_paragraph()
    add_paragraph(doc, 'Estimated effort: 1-2 weeks', italic=True)

    doc.add_page_break()

    add_heading(doc, '3.2 Enhanced Document Support', 2)

    add_paragraph(doc, '3.2.1 PDF with Images/Tables Evaluation', bold=True)
    add_paragraph(doc, 'Current State: Only text extraction from PDFs')
    add_paragraph(doc, 'Improvement: Support PDF images and Claude Vision')

    doc.add_paragraph()
    add_paragraph(doc, 'Use Cases:', bold=True)
    add_bullet(doc, 'Scanned lab reports with handwritten notes')
    add_bullet(doc, 'X-ray/CT images embedded in reports')
    add_bullet(doc, 'Tables with complex lab values')
    add_bullet(doc, 'Charts and graphs')

    doc.add_paragraph()
    add_paragraph(doc, 'Implementation:', bold=True)
    doc.add_paragraph('''1. Use Claude Vision API to analyze images:
   • Convert PDF pages to images
   • Send images to Claude with vision model
   • Extract text and interpret visual elements

2. Add eval test cases with:
   • PDF lab reports with tables
   • Scanned imaging reports
   • Handwritten clinical notes

3. Measure accuracy on visual documents''')

    doc.add_paragraph()
    add_paragraph(doc, 'Technology:', bold=True)
    add_bullet(doc, 'Claude Sonnet 4.5 (has vision capabilities)')
    add_bullet(doc, 'pdf2image library (convert PDF to images)')
    add_bullet(doc, 'Base64 encoding for image upload')

    doc.add_paragraph()
    add_paragraph(doc, 'Estimated effort: 1 week', italic=True)

    doc.add_page_break()

    add_paragraph(doc, '3.2.2 Support More Document Types', bold=True)
    add_paragraph(doc, 'Current State: 5 document types')
    add_paragraph(doc, 'Improvement: Expand to 15+ types')

    doc.add_paragraph()
    add_paragraph(doc, 'New Types to Add:', bold=True)
    add_bullet(doc, 'Lab Reports: Urinalysis, Lipid Panel, Liver Function Tests')
    add_bullet(doc, 'Imaging: MRI, Ultrasound, PET Scan')
    add_bullet(doc, 'Clinical: Progress Notes, Discharge Summary, Consultation')
    add_bullet(doc, 'Administrative: Insurance Forms, Prior Authorization')

    doc.add_paragraph()
    add_paragraph(doc, 'Estimated effort: 2-3 weeks (including eval dataset creation)', italic=True)

    doc.add_page_break()

    add_paragraph(doc, '3.2.3 Image-Only Document Support', bold=True)
    add_paragraph(doc, 'Current State: Text-based documents only')
    add_paragraph(doc, 'Improvement: Upload raw medical images')

    doc.add_paragraph()
    add_paragraph(doc, 'Features:', bold=True)
    add_bullet(doc, 'Upload X-ray/CT/MRI images directly (.jpg, .png, .dicom)')
    add_bullet(doc, 'Claude Vision analyzes image')
    add_bullet(doc, 'Extract findings and ICD-10 codes from image')
    add_bullet(doc, 'Generate radiology report')

    doc.add_paragraph()
    add_paragraph(doc, 'Example Use Cases:', bold=True)
    add_bullet(doc, 'Chest X-ray → Detect pneumonia → Code J18.9')
    add_bullet(doc, 'CT Head → Identify stroke → Code I63.9')
    add_bullet(doc, 'Knee X-ray → Diagnose arthritis → Code M17.9')

    doc.add_paragraph()
    add_paragraph(doc, 'Implementation:', bold=True)
    doc.add_paragraph('''1. Accept image uploads in frontend
2. Send image to Claude Vision API
3. Prompt: "Analyze this medical image. Provide:
   - Modality (X-ray, CT, etc.)
   - Body part examined
   - Key findings
   - Impression
   - Suggested ICD-10 codes"
4. Display results in UI''')

    doc.add_paragraph()
    add_paragraph(doc, 'Estimated effort: 3-5 days', italic=True)

    doc.add_page_break()

    add_heading(doc, '3.3 Database and Storage', 2)

    add_paragraph(doc, '3.3.1 Migrate from SQLite to PostgreSQL/RDS', bold=True)
    add_paragraph(doc, 'Current State: Ephemeral SQLite database')
    add_paragraph(doc, 'Improvement: Persistent PostgreSQL on AWS RDS')

    doc.add_paragraph()
    add_paragraph(doc, 'Why Migrate?', bold=True)
    add_bullet(doc, 'SQLite is ephemeral (data lost when pod restarts)')
    add_bullet(doc, 'No support for concurrent writes')
    add_bullet(doc, 'Not suitable for production')

    doc.add_paragraph()
    add_paragraph(doc, 'Benefits of PostgreSQL:', bold=True)
    add_bullet(doc, 'Persistent data storage')
    add_bullet(doc, 'Automatic backups')
    add_bullet(doc, 'Point-in-time recovery')
    add_bullet(doc, 'Read replicas for scalability')
    add_bullet(doc, 'Better performance at scale')

    doc.add_paragraph()
    add_paragraph(doc, 'AWS RDS Features:', bold=True)
    add_bullet(doc, 'Managed database (no manual maintenance)')
    add_bullet(doc, 'Automatic patches and updates')
    add_bullet(doc, 'Multi-AZ deployment for high availability')
    add_bullet(doc, 'Encryption at rest')

    doc.add_paragraph()
    add_paragraph(doc, 'Migration Steps:', bold=True)
    doc.add_paragraph('''1. Create RDS PostgreSQL instance (db.t3.micro)
2. Update backend DB_URL to postgres://...
3. Run database migrations
4. Update k8s deployment to use RDS
5. Test thoroughly before deploying''')

    doc.add_paragraph()
    add_paragraph(doc, 'Cost Impact: ~$15-20/month for db.t3.micro', italic=True)
    add_paragraph(doc, 'Estimated effort: 1-2 days', italic=True)

    doc.add_page_break()

    add_paragraph(doc, '3.3.2 Store Uploaded Files in S3', bold=True)
    add_paragraph(doc, 'Current State: Files stored locally (ephemeral)')
    add_paragraph(doc, 'Improvement: Store files in S3')

    doc.add_paragraph()
    add_paragraph(doc, 'Benefits:', bold=True)
    add_bullet(doc, 'Persistent file storage')
    add_bullet(doc, 'Unlimited scalability')
    add_bullet(doc, 'High durability (99.999999999%)')
    add_bullet(doc, 'Lifecycle policies for cost optimization')
    add_bullet(doc, 'Versioning for file history')

    doc.add_paragraph()
    add_paragraph(doc, 'Implementation:', bold=True)
    doc.add_paragraph('''Already partially implemented!
• Backend has s3_storage.py module
• Just needs to be enabled:
  STORAGE_BACKEND=s3 (instead of local)
• Files stored at: s3://med-docs-dev/documents/{doc_id}/file.pdf''')

    doc.add_paragraph()
    add_paragraph(doc, 'Cost Impact: ~$0.023 per GB/month', italic=True)
    add_paragraph(doc, 'Estimated effort: Already implemented, just enable!', italic=True)

    doc.add_page_break()

    add_heading(doc, '3.4 Evaluation Improvements', 2)

    add_paragraph(doc, '3.4.1 Expand Eval Dataset to 100+ Cases', bold=True)
    add_paragraph(doc, 'Current State: 50 test cases')
    add_paragraph(doc, 'Target: 100+ cases with edge cases')

    doc.add_paragraph()
    add_paragraph(doc, 'Coverage Goals:', bold=True)
    add_bullet(doc, '20 cases per document type (5 types × 20 = 100)')
    add_bullet(doc, 'Include edge cases:')
    doc.add_paragraph('  • Documents with multiple diagnoses')
    doc.add_paragraph('  • Ambiguous findings')
    doc.add_paragraph('  • Normal results (negative cases)')
    doc.add_paragraph('  • Complex multi-system reports')

    doc.add_paragraph()
    add_paragraph(doc, 'Data Sources:', bold=True)
    add_bullet(doc, 'Synthetic data generated by Claude')
    add_bullet(doc, 'De-identified real reports (HIPAA compliant)')
    add_bullet(doc, 'Public medical datasets (MIMIC-III)')

    doc.add_paragraph()
    add_paragraph(doc, 'Estimated effort: 1 week', italic=True)

    doc.add_page_break()

    add_paragraph(doc, '3.4.2 Add Automated Regression Testing', bold=True)
    add_paragraph(doc, 'Current State: Manual eval review')
    add_paragraph(doc, 'Improvement: Block deployments if metrics drop')

    doc.add_paragraph()
    add_paragraph(doc, 'Implementation:', bold=True)
    doc.add_paragraph('''1. Set minimum thresholds in CI/CD:
   • F1 Score: Must be ≥ 65%
   • Recall: Must be ≥ 80%
   • Precision: Must be ≥ 55%

2. GitHub Actions step:
   if eval_f1 < 0.65:
       print("❌ Eval metrics dropped below threshold!")
       exit(1)

3. Prevent bad code from reaching production''')

    doc.add_paragraph()
    add_paragraph(doc, 'Estimated effort: 2-3 hours', italic=True)

    doc.add_page_break()

    add_heading(doc, '3.5 User Experience', 2)

    add_paragraph(doc, '3.5.1 Batch Document Upload', bold=True)
    add_paragraph(doc, 'Current State: Upload one file at a time')
    add_paragraph(doc, 'Improvement: Upload multiple files simultaneously')

    doc.add_paragraph()
    add_paragraph(doc, 'Features:', bold=True)
    add_bullet(doc, 'Drag-and-drop multiple files')
    add_bullet(doc, 'Apply same document type hint to all')
    add_bullet(doc, 'Progress bar for batch processing')
    add_bullet(doc, 'Summary of results in table format')

    doc.add_paragraph()
    add_paragraph(doc, 'Estimated effort: 3-5 days', italic=True)

    doc.add_paragraph()
    add_paragraph(doc, '3.5.2 Export Results to CSV/Excel', bold=True)
    add_paragraph(doc, 'Current State: View results in UI only')
    add_paragraph(doc, 'Improvement: Export to spreadsheet')

    doc.add_paragraph()
    add_paragraph(doc, 'Features:', bold=True)
    add_bullet(doc, 'Export button on results page')
    add_bullet(doc, 'Include: document name, type, ICD-10 codes, confidence scores')
    add_bullet(doc, 'Support CSV and Excel formats')
    add_bullet(doc, 'Useful for billing departments')

    doc.add_paragraph()
    add_paragraph(doc, 'Estimated effort: 1-2 days', italic=True)

    doc.add_page_break()

    add_paragraph(doc, '3.5.3 Real-Time Processing Status', bold=True)
    add_paragraph(doc, 'Current State: Loading spinner only')
    add_paragraph(doc, 'Improvement: Show progress of each pipeline stage')

    doc.add_paragraph()
    add_paragraph(doc, 'Features:', bold=True)
    add_bullet(doc, 'Progress indicator with stages:')
    doc.add_paragraph('  ✓ Uploading file...')
    doc.add_paragraph('  ✓ Extracting text...')
    doc.add_paragraph('  🔄 Classifying document... (5s)')
    doc.add_paragraph('  ⏳ Extracting codes... (waiting)')
    add_bullet(doc, 'WebSocket connection for real-time updates')
    add_bullet(doc, 'Estimated completion time')

    doc.add_paragraph()
    add_paragraph(doc, 'Estimated effort: 3-5 days', italic=True)

    doc.add_page_break()

    add_heading(doc, '3.6 Monitoring and Observability', 2)

    add_paragraph(doc, '3.6.1 Application Performance Monitoring (APM)', bold=True)
    add_paragraph(doc, 'Current State: Basic CloudWatch logs')
    add_paragraph(doc, 'Improvement: Comprehensive APM with DataDog/New Relic')

    doc.add_paragraph()
    add_paragraph(doc, 'Features:', bold=True)
    add_bullet(doc, 'Request tracing (see full lifecycle of each request)')
    add_bullet(doc, 'Performance metrics (API latency, throughput)')
    add_bullet(doc, 'Error tracking with stack traces')
    add_bullet(doc, 'Custom dashboards')
    add_bullet(doc, 'Alerting (Slack/email notifications)')

    doc.add_paragraph()
    add_paragraph(doc, 'Recommended Tool: DataDog', bold=True)
    add_bullet(doc, 'Easy Kubernetes integration')
    add_bullet(doc, 'Python APM agent')
    add_bullet(doc, 'Free tier available')

    doc.add_paragraph()
    add_paragraph(doc, 'Estimated effort: 1 day setup + ongoing monitoring', italic=True)

    doc.add_paragraph()
    add_paragraph(doc, '3.6.2 Cost Tracking Dashboard', bold=True)
    add_paragraph(doc, 'Current State: Manual cost review in AWS Console')
    add_paragraph(doc, 'Improvement: Real-time cost dashboard')

    doc.add_paragraph()
    add_paragraph(doc, 'Metrics to Track:', bold=True)
    add_bullet(doc, 'Claude API usage and costs')
    add_bullet(doc, 'AWS service costs (EKS, S3, ECR, etc.)')
    add_bullet(doc, 'Cost per document processed')
    add_bullet(doc, 'Monthly burn rate')
    add_bullet(doc, 'Cost anomaly detection')

    doc.add_paragraph()
    add_paragraph(doc, 'Estimated effort: 2-3 days', italic=True)

    doc.add_page_break()

    add_heading(doc, '3.7 Additional Features', 2)

    add_paragraph(doc, '3.7.1 Confidence Threshold Settings', bold=True)
    add_paragraph(doc, 'Allow users to set minimum confidence for ICD-10 codes')
    add_bullet(doc, 'Default: 0.5')
    add_bullet(doc, 'High confidence: 0.8+ (fewer codes, more accurate)')
    add_bullet(doc, 'Low confidence: 0.3+ (more codes, may include false positives)')

    doc.add_paragraph()
    add_paragraph(doc, '3.7.2 Custom Prompt Templates', bold=True)
    add_paragraph(doc, 'Let admins customize AI prompts for specific use cases')
    add_bullet(doc, 'Specialty-specific templates (cardiology, oncology)')
    add_bullet(doc, 'Language preferences')
    add_bullet(doc, 'Detail level (brief vs comprehensive summaries)')

    doc.add_paragraph()
    add_paragraph(doc, '3.7.3 Multi-Language Support', bold=True)
    add_paragraph(doc, 'Support medical documents in multiple languages')
    add_bullet(doc, 'Spanish, French, German, Hindi')
    add_bullet(doc, 'Auto-detect language')
    add_bullet(doc, 'Translate results to English')

    doc.add_paragraph()
    add_paragraph(doc, '3.7.4 Integration with EHR Systems', bold=True)
    add_paragraph(doc, 'Connect with Electronic Health Record systems')
    add_bullet(doc, 'FHIR API support')
    add_bullet(doc, 'HL7 message parsing')
    add_bullet(doc, 'Epic/Cerner integration')

    doc.add_page_break()

    # Priority Roadmap
    add_heading(doc, '4. Priority Roadmap (Next 3 Months)', 1)

    add_paragraph(doc, 'Month 1: Security & Stability', bold=True)
    add_bullet(doc, '✓ Move API key to AWS Secrets Manager (PRIORITY 1)')
    add_bullet(doc, '✓ Migrate to PostgreSQL/RDS')
    add_bullet(doc, '✓ Enable S3 file storage')
    add_bullet(doc, '✓ Set up DataDog monitoring')

    doc.add_paragraph()
    add_paragraph(doc, 'Month 2: Enhanced Document Support', bold=True)
    add_bullet(doc, '✓ Add PDF image/table evaluation')
    add_bullet(doc, '✓ Support image-only documents')
    add_bullet(doc, '✓ Expand eval dataset to 100+ cases')
    add_bullet(doc, '✓ Add regression testing gates')

    doc.add_paragraph()
    add_paragraph(doc, 'Month 3: User Experience', bold=True)
    add_bullet(doc, '✓ Batch upload feature')
    add_bullet(doc, '✓ Export to CSV/Excel')
    add_bullet(doc, '✓ Real-time processing status')
    add_bullet(doc, '✓ User authentication (Cognito)')

    doc.add_paragraph()
    add_heading(doc, 'Long-Term (6-12 Months)', 2)
    add_bullet(doc, 'Support 15+ document types')
    add_bullet(doc, 'Multi-language support')
    add_bullet(doc, 'EHR system integration (FHIR)')
    add_bullet(doc, 'Advanced analytics and reporting')
    add_bullet(doc, 'Mobile app (iOS/Android)')

    doc.add_page_break()

    # Summary
    add_heading(doc, '5. Summary', 1)

    add_paragraph(doc, 'MedicalDocAI has a solid foundation with core features working well:', bold=True)

    doc.add_paragraph()
    add_bullet(doc, '5 document types with high accuracy')
    add_bullet(doc, 'ICD-10 code extraction with evidence')
    add_bullet(doc, 'Clinical summaries for providers')
    add_bullet(doc, 'Document type hints to improve speed')
    add_bullet(doc, 'Comprehensive evaluation system')
    add_bullet(doc, 'Production-ready cloud deployment')

    doc.add_paragraph()
    add_paragraph(doc, 'Key priorities for future development:', bold=True)

    doc.add_paragraph()
    add_paragraph(doc, '1. Security (API key in Secrets Manager)', bold=True)
    doc.add_paragraph('   Critical for production readiness and compliance')

    doc.add_paragraph()
    add_paragraph(doc, '2. Enhanced Document Support (PDF images, medical images)', bold=True)
    doc.add_paragraph('   Expands use cases significantly')

    doc.add_paragraph()
    add_paragraph(doc, '3. User Authentication and Access Control', bold=True)
    doc.add_paragraph('   Required for HIPAA compliance and multi-tenant usage')

    doc.add_paragraph()
    add_paragraph(doc, '4. Persistent Storage (PostgreSQL + S3)', bold=True)
    doc.add_paragraph('   Essential for production-grade reliability')

    doc.add_paragraph()
    doc.add_paragraph('With these improvements, MedicalDocAI can become a enterprise-ready medical document processing platform serving hospitals, clinics, and healthcare organizations.')

    # Save document
    doc_path = os.path.join(os.path.dirname(__file__), '..', 'docs', 'Features_and_Roadmap.docx')
    os.makedirs(os.path.dirname(doc_path), exist_ok=True)
    doc.save(doc_path)
    print(f"[OK] Created: {doc_path}")
    return doc_path


if __name__ == "__main__":
    print("Generating project documentation...\n")

    # Create both documents
    infra_doc = create_infrastructure_doc()
    features_doc = create_features_roadmap_doc()

    print(f"\nDocumentation generated successfully!")
    print(f"\nDocuments saved to:")
    print(f"  - {infra_doc}")
    print(f"  - {features_doc}")
    print(f"\nYou can find them in the docs/ folder")
